
import os
from docx import *
###需要新增 docx 文件的读取

####表格cells文字的读取测试，打印全部cell的信息
def opDocx():
    doc = Document("test.docx")

    for i, table in enumerate(doc.tables):
        for j, row in enumerate(table.rows):
            record = ""
            for k, cell in enumerate(row.cells):
                if k == 0:
                    print("第{}个table 第{}行 第{}个 cell ==> {}".format(i, j, k , cell.text))
                    record = row.cells[k].text
                elif row.cells[k].text != record:
                    print("第{}个table 第{}行 第{}个 cell ==> {}".format(i, j, k , cell.text))
                else:
                    print("第{}个table {}行 第{}个 cell 同上".format(i, j, k))

####表格字体设置的测试
def opTem():
    doc = Document("test.docx")
    table = doc.tables[0]
    styleTemp = table.rows[11].cells[2].paragraphs[0].runs[0]
    #for i in range(1,2):
    run = table.rows[1].cells[2].paragraphs[0].add_run('smida')
        #.paragraphs[0].add_run("123",styleTemp)
    run.font.name = '仿宋_GB2312'
    run.font.size = 180000
    #run.font = styleTemp.font
    #160000 ->12.5   170000 -> 13   175000 -> 13.5
    doc.save("test2.docx")


def main():
    with open("test.apk", "rb") as f:
        rawFile = f.read()

    if rawFile is None:
        print("打开apk文件失败")
        f.close()
        return
    global targeApk
    #### 用列表搜集需要的信息
    cell_List = []
    targeApk = apk.APK(rawFile,True)
    targeDex = targeApk.get_dex()
    apkMd5 = hashlib.md5(rawFile).hexdigest().upper()
    ##文件名称、文件MD5、DEX MD5、文件大小
    fileName = apkMd5 + ".apk"
    cell_List.append(fileName)
    cell_List.append(apkMd5)
    dexMd5 = hashlib.md5(targeDex).hexdigest().upper()
    cell_List.append(dexMd5)
    fileSize = format(os.path.getsize(f.name),',') + " 字节"
    cell_List.append(fileSize.lstrip())
    ##app名称、包名、签名（issuer）、证书串号、公钥MD5、公钥SHA1、
    apkName = targeApk.get_app_name()
    cell_List.append(apkName)
    apkPackageName = targeApk.get_package()
    cell_List.append(apkPackageName)
    apkSign  = getSign()  ###一段蹩脚的签名获取方法，后去可用keytool工具直接获取
    cell_List.append(apkSign["Issuer"])
    cell_List.append(apkSign["SerialNumber"])
    cell_List.append(apkSign["signMd5"])
    cell_List.append(apkSign["signSha1"])

    ###  操作word
    doc = Document("test.docx")
    table = doc.tables[0]
    #table.rows[1].cells[2].text = fileName    这种方式废弃
    #table.rows[2].cells[2].text = apkMd5
    #table.rows[3].cells[2].text = dexMd5


    for i in range(1, 11):
        run = table.rows[i].cells[2].paragraphs[0].add_run(cell_List[i-1])  ###或者写成  table.rows(i, 2)
        run.font.name = '仿宋_GB2312'
        run.font.size = 175000

    doc.save("test1.docx")
    ###是否结束？清理工作
    print("everything is done, exit?")
    _ = input()
    f.close()
    return