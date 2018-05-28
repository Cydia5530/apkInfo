from docx import Document

##### C 语言版
def printDocxAsC():
    doc = Document("test.docx")
    tableNums = len(doc.tables)
    record = ""
    for k in range(tableNums):
        table = doc.tables[k]
        rowNums = len(table.rows)
        for i in range(rowNums):
            row = table.rows[i]
            cellNums = len(row.cells)
            for j in range(cellNums):
                if j == 0:
                    print("第{}个table的 第{}行 第{}个 cell ==> {}".format(k + 1, i + 1, j + 1, row.cells[j].text))
                    record = row.cells[j].text
                elif row.cells[j].text != record:
                    print("第{}个table的 第{}行 第 {}个cell ==> {}".format(k + 1, i + 1, j + 1, row.cells[j].text))
                else:
                    print("第{}个table的 第{}行 第 {}个cell 同上".format(k + 1, i + 1, j + 1))

##### python 版
def printDocxAsP():
    doc = Document("test.docx")

    for i, table in enumerate(doc.tables):
        for j, row in enumerate(table.rows):
            record = ""
            for k, cell in enumerate(row.cells):
                if k == 0:
                    print("第{}个table 第{}行 第{}个 cell ==> {}".format(i + 1, j + 1, k + 1, cell.text))
                    record = row.cells[k].text
                elif row.cells[k].text != record:
                    print("第{}个table 第{}行 第{}个 cell ==> {}".format(i + 1, j + 1, k + 1, cell.text))
                else:
                    print("第{}个table {}行 第{}个 cell 同上".format(i + 1, j + 1, k + 1))


if __name__ == '__main__':
    printDocxAsP()





